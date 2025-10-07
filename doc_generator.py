# doc_generator.py

import docx
import io
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_section_header(doc, title):
    """Adds a styled section header with a horizontal line."""
    p = doc.add_paragraph()
    p.add_run(title.upper()).bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    
    # Add a horizontal line (bottom border) to the paragraph
    p_bdr = OxmlElement('w:pBdr')
    bottom_bdr = OxmlElement('w:bottom')
    bottom_bdr.set(qn('w:val'), 'single')
    bottom_bdr.set(qn('w:sz'), '6')
    bottom_bdr.set(qn('w:space'), '1')
    bottom_bdr.set(qn('w:color'), 'auto')
    p_bdr.append(bottom_bdr)
    p._p.get_or_add_pPr().append(p_bdr)

def create_professional_template(full_name, email, phone_number, linkedin_profile, summary, experience, education, skills, certifications):
    """Creates a stylish, template-based .docx file."""
    doc = docx.Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # --- HEADER ---
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(full_name.upper())
    name_run.bold = True
    name_run.font.size = Pt(24)

    # --- This is the corrected line ---
    #tagline_p = doc.add_paragraph("Digital Marketing | SEO | SEM | Content Marketing")
    #tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.add_run(f"Email: {email} | Phone Number: {phone_number}\n LinkedIn Profile: {linkedin_profile}") # Added location placeholder
    contact_p.paragraph_format.space_after = Pt(18)


    # --- PROFESSIONAL SUMMARY ---
    add_section_header(doc, "Professional Summary")
    doc.add_paragraph(summary)

    # --- WORK EXPERIENCE ---
    add_section_header(doc, "Work Experience")
    doc.add_paragraph(experience)


    # --- EDUCATION ---
    add_section_header(doc, "Education")
    doc.add_paragraph(education)
    

    # --- SKILLS ---
    add_section_header(doc, "Skills")
    doc.add_paragraph(skills)


    # --- CERTIFICATIONS ---
    add_section_header(doc, "Certifications")
    doc.add_paragraph(certifications)
    
    # Save to a byte stream to be used by Streamlit
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
# doc_generator.py

def add_modern_section_header(paragraph, title, draw_line=True):
    """Helper for the Modern Template's section headers."""
    if draw_line:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        top_bdr = OxmlElement('w:top')
        top_bdr.set(qn('w:val'), 'single'); top_bdr.set(qn('w:sz'), '4')
        p_bdr.append(top_bdr)
        p_pr.append(p_bdr)
        # Add space above the line
        paragraph.paragraph_format.space_before = Pt(12)
    
    # Add a newline character before the title to create the gap
    run = paragraph.add_run(f"\n{title.upper()}")
    run.bold = True
    run.font.size = Pt(11)
    paragraph.paragraph_format.space_after = Pt(6)


def create_modern_template(full_name, email, phone_number, linkedin_profile, summary, experience, education, skills, certifications):
    doc = docx.Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.7); section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75); section.right_margin = Inches(0.75)
    
    # Set default paragraph spacing
    style = doc.styles['Normal']
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(4)
    paragraph_format.line_spacing = 1.15

    # Header
    name_p = doc.add_paragraph(); name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(full_name.upper()); name_run.font.size = Pt(24)
    name_run.bold = True
    name_p.paragraph_format.space_after = Pt(10)

    # --- Main two-column layout using a styled table ---
    table = doc.add_table(rows=1, cols=2)
    # (Code to style the table with a vertical line is unchanged)
    tbl_pr = table._element.xpath('w:tblPr')[0]
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH']:
        border = OxmlElement(f'w:{border_name}'); border.set(qn('w:val'), 'nil')
        tbl_borders.append(border)
    inside_v = OxmlElement('w:insideV'); inside_v.set(qn('w:val'), 'single'); inside_v.set(qn('w:sz'), '4')
    tbl_borders.append(inside_v)
    tbl_pr.append(tbl_borders)

    left_cell = table.cell(0, 0); right_cell = table.cell(0, 1)
    left_cell.width = Inches(2.8); right_cell.width = Inches(4.2)
    left_cell.paragraphs[0].paragraph_format.right_indent = Inches(0.2)
    right_cell.paragraphs[0].paragraph_format.left_indent = Inches(0.2)

    # --- Populate Left & Right Columns ---
    
    # Right-side headers get an indent
    summary_header_p = right_cell.add_paragraph()
    summary_header_p.paragraph_format.left_indent = Inches(0.25)
    add_modern_section_header(summary_header_p, "Profile Summary", draw_line=False)
    summary_p = right_cell.add_paragraph(summary); summary_p.paragraph_format.left_indent = Inches(0.25)
    
    experience_header_p = right_cell.add_paragraph()
    experience_header_p.paragraph_format.left_indent = Inches(0.25)
    add_modern_section_header(experience_header_p, "Work Experience")
    experience_p = right_cell.add_paragraph(experience); experience_p.paragraph_format.left_indent = Inches(0.25)

    # Left-side headers do not
    add_modern_section_header(left_cell.add_paragraph(), "Contact", draw_line=False)
    left_cell.add_paragraph(phone_number); left_cell.add_paragraph(email); left_cell.add_paragraph(linkedin_profile)
    add_modern_section_header(left_cell.add_paragraph(), "Education"); left_cell.add_paragraph(education)
    add_modern_section_header(left_cell.add_paragraph(), "Skills"); left_cell.add_paragraph(skills)
    add_modern_section_header(left_cell.add_paragraph(), "Certifications"); left_cell.add_paragraph(certifications)
    
    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio